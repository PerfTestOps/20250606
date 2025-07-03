import streamlit as st
import pandas as pd
from docx import Document
from fpdf import FPDF
import plotly.express as px
import matplotlib.pyplot as plt
from docx.shared import Inches
from datetime import datetime
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def transaction_analysis_page():

    # st.set_page_config(page_title="Transaction Analyzer", page_icon=":bar_chart:")

    # Define a function for generating a pie chart for different runs
    def generatepiechart(filtered_df1):
        num_rowsPie, num_colsPie = filtered_df1.shape
        for i in range(1, num_colsPie):
            if i < num_colsPie:                 
                sla_column = filtered_df1.columns[1]  # SLA Values column
                current_column = filtered_df1.columns[i]
                sla_breach_count = 0
                if i >= 2:  # ensure reading the runs column
                    for value, SLAValue in zip(filtered_df1[current_column], filtered_df1[sla_column]):
                        if value > SLAValue:
                            sla_breach_count += 1
                    labels = ['SLA Met', 'SLA Breached']
                    sizes = [num_rowsPie - sla_breach_count, sla_breach_count]
                    colors = ['green', 'red']
                    explode = (0.1, 0)
                    fig, ax = plt.subplots(figsize=(0.8, 0.8), dpi=1000)
                    ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', shadow=True, startangle=140, textprops={'fontsize': 5})
                    plt.title(f"{current_column}", fontsize=6)
                    st.pyplot(fig)

    # Define a function to apply styling
    def highlight_sla(row, sla_col='SLA', tolerance=0):
        sla = row[sla_col]
        result = []
        for col, value in row.items():
            if col == sla_col or col == 'TransactionName':
                result.append('')
            else:
                try:
                    if float(value) > (1 + tolerance / 100) * float(sla):
                        result.append('background-color: red')
                    else:
                        result.append('background-color: green')
                except:
                    result.append('')
        return result

    # Generate a report
    def generate_report(filtered_df1, graph_path, tolerance=0):
        doc = Document()
        doc.add_heading("Performance Test Report", level=1)

        # Add DataFrame as a Table
        table = doc.add_table(rows=1, cols=len(filtered_df1.columns))
        table.style = "Table Grid"

        # Header Row
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(filtered_df1.columns):
            hdr_cells[j].text = col

        # Data Rows with Highlighting
        for i, row in filtered_df1.iterrows():
            row_cells = table.add_row().cells
            sla = row.get('SLA', None)

            for j, col in enumerate(filtered_df1.columns):
                value = row[col]
                cell = row_cells[j]
                cell.text = str(value)

                # Apply background color only to run columns (exclude TransactionName and SLA)
                if col not in ['TransactionName', 'SLA']:
                    try:
                        if float(value) > (1 + tolerance / 100) * float(sla):
                            shade = "FF0000"  # red
                        else:
                            shade = "00FF00"  # green

                        # Add shading (background color)
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), shade)
                        tcPr.append(shd)

                    except:
                        pass  # skip if conversion fails

        # Add chart
        doc.add_heading("Performance Chart", level=1)
        doc.add_picture(graph_path, width=Inches(5))

        # Save document
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"PerformanceTestReport_{timestamp}.docx"
        doc.save(filename)
        st.success(f"Word document saved successfully as: {filename}")

    # Load Excel file
    uploaded_file = st.file_uploader("Choose a file", type=["csv", "xlsx", "txt"])

    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file, index_col=None)

        st.title("Data Query and Graph Generator")

        with st.expander("View DataFrame"):
            st.dataframe(df)

        columns = df.columns.tolist()
        column_to_filter = st.selectbox("Select column to filter", columns)
        unique_values = df[column_to_filter].unique()
        filter_value = st.selectbox("Select value to filter by", unique_values)
        filtered_df = df[df[column_to_filter] == filter_value]

        with st.expander("View Filtered DataFrame"):
            st.dataframe(filtered_df)

        del filtered_df['Status']

        st.sidebar.header("Please Filter Here:")
        selected_columns = st.sidebar.multiselect("Select columns to include", filtered_df.columns.tolist(), default=filtered_df.columns.tolist())

        if selected_columns:
            filtered_df1 = filtered_df[selected_columns]

            st.write("Filtered Table:")
            show_message = st.checkbox('Highlight SLA Deviations')

            # SLA Tolerance and Highlighting Logic
            tolerance = st.sidebar.selectbox("SLA Tolerance Percentage", [0, 10, 20, 30, 40, 50])

            if show_message:
                styled_df = filtered_df1.style.apply(lambda row: highlight_sla(row, sla_col='SLA', tolerance=tolerance), axis=1)
                st.dataframe(styled_df)
            else:
                st.dataframe(filtered_df1)

            # 📊 Dynamic Graph Builder Section for Response Time
            st.subheader("📊 Response Time Comparison Graph")
            x_axis = st.selectbox("Select X-axis", filtered_df1.columns.tolist(), index=0)
            y_axis = st.multiselect("Select Y-axis Columns", [col for col in filtered_df1.columns if col != x_axis], default=filtered_df1.columns[1:2])

            if x_axis and y_axis:
                graph_title = f"{', '.join(y_axis)} vs {x_axis}"
                fig_dynamic = px.bar(filtered_df1, x=x_axis, y=y_axis, title=graph_title, barmode='group', text_auto=True)
                fig_dynamic.update_layout(xaxis_title=x_axis, yaxis_title="Values")
                st.plotly_chart(fig_dynamic, use_container_width=True)


            # 📊 Dynamic Graph Builder Section for TPH
            st.subheader("📊 TPH Comparison Graph")
            x_axis = st.selectbox("Select X-axis for TPH", filtered_df1.columns.tolist(), index=0)
            y_axis = st.multiselect("Select Y-axis Columns for TPH", [col for col in filtered_df1.columns if col != x_axis], default=filtered_df1.columns[2:3])

            if x_axis and y_axis:
                graph_title = f"{', '.join(y_axis)} vs {x_axis}"
                fig_dynamicTPH = px.bar(filtered_df1, x=x_axis, y=y_axis, title=graph_title, barmode='group', text_auto=True)
                fig_dynamicTPH.update_layout(xaxis_title=x_axis, yaxis_title="Values")
                st.plotly_chart(fig_dynamicTPH, use_container_width=True)

            # 📊 Dynamic Graph Builder Section for Error %
            st.subheader("📊 Error Comparison Graph")
            x_axis = st.selectbox("Select X-axis for Error%", filtered_df1.columns.tolist(), index=0)
            y_axis = st.multiselect("Select Y-axis Columns for Error%", [col for col in filtered_df1.columns if col != x_axis])

            if x_axis and y_axis:
                graph_title = f"{', '.join(y_axis)} vs {x_axis}"
                fig_dynamicError = px.bar(filtered_df1, x=x_axis, y=y_axis, title=graph_title, barmode='group', text_auto=True)
                fig_dynamicError.update_layout(xaxis_title=x_axis, yaxis_title="Values")
                st.plotly_chart(fig_dynamicError, use_container_width=True)

        st.title(":bar_chart: Transaction Analysis")
        st.bar_chart(filtered_df1.set_index('TransactionName'), stack=False)

        fig, ax = plt.subplots(figsize=(10, 6))
        ax.plot(filtered_df1['TransactionName'], filtered_df1['SLA'], marker='o', label='SLA', color='red', linewidth=2)
        bar_width = 0.2
        x_indexes = range(len(filtered_df1['TransactionName']))

        for i, column in enumerate(filtered_df1.columns[2:]):
            ax.bar([x + i * bar_width for x in x_indexes], filtered_df1[column], bar_width, label=column)

        ax.set_xlabel('Transaction Names')
        ax.set_ylabel('90 Percent Response Times (secs)')
        ax.set_title('Bar Chart with Fixed X-axis and Varying Y-columns')
        ax.set_xticks([x + bar_width for x in x_indexes])
        ax.set_xticklabels(filtered_df1['TransactionName'])
        ax.legend()
        ax.grid(axis='y', linestyle='--', alpha=0.7)
        plt.savefig('plot.png', format='png')
        st.pyplot(fig)

        graph_path = "plot.png"
        fig.savefig(graph_path)

        # Generate pie chart
        generatepiechart(filtered_df1)

        if st.button('Generate Report'):
        #generate_report(filtered_df1) 
            #st.bar_chart(filtered_df1.set_index('TransactionName'), stack=False) 
            
            generate_report(filtered_df1, graph_path, tolerance) 
        #fig = px.bar(df, x='TransactionName',title='Bar Chart Example')
        #fig.show()
    else:
        st.write("")
