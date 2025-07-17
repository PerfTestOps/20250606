import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
from datetime import datetime
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import matplotlib.pyplot as plt

def transaction_analysis_page():
    def generatepiechart(filtered_df1):
        num_rowsPie, num_colsPie = filtered_df1.shape
        for i in range(2, num_colsPie):  # Start from 2nd column after SLA
            sla_column = filtered_df1.columns[1]
            current_column = filtered_df1.columns[i]
            sla_breach_count = 0
            for value, SLAValue in zip(filtered_df1[current_column], filtered_df1[sla_column]):
                if value > SLAValue:
                    sla_breach_count += 1
            labels = ['SLA Met', 'SLA Breached']
            sizes = [num_rowsPie - sla_breach_count, sla_breach_count]
            colors = ['green', 'red']
            explode = (0.1, 0)
            fig, ax = plt.subplots(figsize=(2, 2), dpi=200)
            ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%',
                   shadow=True, startangle=140, textprops={'fontsize': 6})
            plt.title(f"{current_column}", fontsize=8)
            st.pyplot(fig)

    def highlight_sla(row):
        sla = row['SLA']
        tolerance_pct = st.session_state.get('tolerance', 0)
        styles = []
        for col in row.index:
            if col.endswith('Execution'):
                try:
                    value = float(row[col])
                    if value <= sla * (1 + tolerance_pct / 100):
                        styles.append('background-color: lightgreen')
                    else:
                        styles.append('background-color: lightcoral')
                except:
                    styles.append('')
            else:
                styles.append('')
        return styles

    def generate_report(filtered_df1, tolerance=0):
        doc = Document()
        doc.add_heading("Performance Test Report", level=1)

        # Add data table
        table = doc.add_table(rows=1, cols=len(filtered_df1.columns))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(filtered_df1.columns):
            hdr_cells[j].text = col
        for i, row in filtered_df1.iterrows():
            row_cells = table.add_row().cells
            sla_value = float(row['SLA']) if 'SLA' in row else None
            for j, (col_name, value) in enumerate(row.items()):
                cell = row_cells[j]
                cell.text = str(value)
                if sla_value is not None and 'Execution' in col_name:
                    try:
                        if float(value) > (1 + tolerance / 100) * float(sla_value):
                            shade = "FF0000"
                        else:
                            shade = "00FF00"
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), shade)
                        tcPr.append(shd)
                    except:
                        pass

        # Add charts
        doc.add_heading("Performance Charts", level=1)
        for title, filename in [
            ("Response Time Comparison Graph", "response_time_chart.png"),
            ("TPH Comparison Graph", "tph_chart.png"),
            ("Error Percentage Comparison Graph", "error_chart.png"),
            ("SLA Comparison Graph", "final_sla_chart.png")
        ]:
            doc.add_paragraph(title)
            doc.add_picture(filename, width=Inches(5))

        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"PerformanceTestReport_{timestamp}.docx"
        doc.save(filename)
        st.success(f"Word document saved successfully as: {filename}")

    uploaded_file = st.file_uploader("Choose a file", type=["csv", "xlsx", "txt"])

    if uploaded_file is not None:
        df = pd.read_excel(uploaded_file, index_col=None)
        st.title("Data Query and Graph Generator")

        with st.expander("View DataFrame"):
            st.dataframe(df)

        columns = df.columns.tolist()
        column_to_filter = st.selectbox("Select column to filter", columns)
        unique_values = df[column_to_filter].unique()
        filter_value = st.selectbox("Select value to filter by", unique_values)
        filtered_df = df[df[column_to_filter] == filter_value]

        with st.expander("View Filtered DataFrame"):
            st.dataframe(filtered_df)

        if 'Status' in filtered_df.columns:
            del filtered_df['Status']

        st.sidebar.header("Please Filter Here:")
        selected_columns = st.sidebar.multiselect("Select columns to include", filtered_df.columns.tolist(), default=filtered_df.columns.tolist())

        if selected_columns:
            filtered_df1 = filtered_df[selected_columns]
            st.write("Filtered Table:")
            show_message = st.checkbox('Highlight SLA Deviations')
            tolerance = st.sidebar.selectbox("SLA Tolerance Percentage", [0, 10, 20, 30, 40, 50])
            st.session_state['tolerance'] = tolerance

            if show_message:
                styled_df = filtered_df1.style.apply(highlight_sla, axis=1)
                st.dataframe(styled_df)
            else:
                st.dataframe(filtered_df1)

            # Response Time Graph (SLA as line)
            st.subheader("Response Time Comparison Graph")
            x_axis = st.selectbox("Select X-axis", filtered_df1.columns.tolist(), index=0)
            y_axis = st.multiselect("Select Y-axis Columns", [col for col in filtered_df1.columns if col != x_axis], default=filtered_df1.columns[1:2])

            if x_axis and y_axis:
                fig, ax = plt.subplots(figsize=(10, 6))
                bar_width = 0.2
                index = range(len(filtered_df1[x_axis]))

                if 'SLA' in filtered_df1.columns:
                    ax.plot(filtered_df1[x_axis], filtered_df1['SLA'], label='SLA', color='red', marker='o', linewidth=2)

                for i, y in enumerate(y_axis):
                    if y != 'SLA':
                        ax.bar([x + i * bar_width for x in index], filtered_df1[y], width=bar_width, label=y)

                ax.set_title(f"{', '.join(y_axis)} vs {x_axis}")
                ax.set_xlabel(x_axis)
                ax.set_ylabel("Values")
                ax.set_xticks([x + bar_width for x in index])
                ax.set_xticklabels(filtered_df1[x_axis], rotation=45)
                ax.legend()
                plt.tight_layout()
                fig.savefig("response_time_chart.png")
                st.pyplot(fig)

            # TPH Graph (Target TPH as line)
            st.subheader("TPH Comparison Graph")
            x_axis = st.selectbox("Select X-axis for TPH", filtered_df1.columns.tolist(), index=0)
            y_axis = st.multiselect("Select Y-axis Columns for TPH", [col for col in filtered_df1.columns if col != x_axis], default=filtered_df1.columns[2:3])

            if x_axis and y_axis:
                fig, ax = plt.subplots(figsize=(10, 6))
                bar_width = 0.2
                index = range(len(filtered_df1[x_axis]))

                if 'Target TPH' in filtered_df1.columns:
                    ax.plot(filtered_df1[x_axis], filtered_df1['Target TPH'], label='Target TPH', color='red', marker='D', linewidth=2)

                for i, y in enumerate(y_axis):
                    if y != 'Target TPH':
                        ax.bar([x + i * bar_width for x in index], filtered_df1[y], width=bar_width, label=y)

                ax.set_title(f"{', '.join(y_axis)} vs {x_axis}")
                ax.set_xlabel(x_axis)
                ax.set_ylabel("Values")
                ax.set_xticks([x + bar_width for x in index])
                ax.set_xticklabels(filtered_df1[x_axis], rotation=45)
                ax.legend()
                plt.tight_layout()
                fig.savefig("tph_chart.png")
                st.pyplot(fig)

            # Error % Graph
            st.subheader("Error Comparison Graph")
            x_axis = st.selectbox("Select X-axis for Error%", filtered_df1.columns.tolist(), index=0)
            y_axis = st.multiselect("Select Y-axis Columns for Error%", [col for col in filtered_df1.columns if col != x_axis])

            if x_axis and y_axis:
                fig, ax = plt.subplots(figsize=(10, 6))
                bar_width = 0.2
                index = range(len(filtered_df1[x_axis]))
                for i, y in enumerate(y_axis):
                    ax.bar([x + i * bar_width for x in index], filtered_df1[y], width=bar_width, label=y)
                ax.set_title(f"{', '.join(y_axis)} vs {x_axis}")
                ax.set_xlabel(x_axis)
                ax.set_ylabel("Values")
                ax.set_xticks([x + bar_width for x in index])
                ax.set_xticklabels(filtered_df1[x_axis], rotation=45)
                ax.legend()
                plt.tight_layout()
                fig.savefig("error_chart.png")
                st.pyplot(fig)

            # SLA Bar Chart (Final Graph)
            st.title(":bar_chart: SLA Comparison Graph")
            fig, ax = plt.subplots(figsize=(10, 6))
            ax.plot(filtered_df1['TransactionName'], filtered_df1['SLA'], marker='o', label='SLA', color='red', linewidth=2)
            bar_width = 0.2
            x_indexes = range(len(filtered_df1['TransactionName']))
            for i, column in enumerate(filtered_df1.columns[2:]):
                ax.bar([x + i * bar_width for x in x_indexes], filtered_df1[column], bar_width, label=column)
            ax.set_xlabel('Transaction Names')
            ax.set_ylabel('90 Percent Response Times (secs)')
            ax.set_title('SLA vs Execution Times')
            ax.set_xticks([x + bar_width for x in x_indexes])
            ax.set_xticklabels(filtered_df1['TransactionName'], rotation=45)
            ax.legend()
            ax.grid(axis='y', linestyle='--', alpha=0.7)
            plt.tight_layout()
            fig.savefig("final_sla_chart.png")
            st.pyplot(fig)

            generatepiechart(filtered_df1)

            if st.button('Generate Report'):
                generate_report(filtered_df1, tolerance)
